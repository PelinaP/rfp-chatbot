import streamlit as st
import os
import tempfile
import shutil
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama
from langchain.chains import RetrievalQA
from docx import Document

st.set_page_config(page_title="RFP Chatbot", layout="wide")
st.title("📄 RFP Chatbot")

# Select LLM model
model_name = st.selectbox("Select LLM Model", ["mistral", "llama2", "llama3"])

# Initialize chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

uploaded_file = st.file_uploader("Upload your RFP PDF", type=["pdf"])

if uploaded_file:
    st.success("✅ RFP uploaded successfully.")
    st.write(f"**Filename:** {uploaded_file.name}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.read())
        temp_path = tmp_file.name

    # Load and embed the PDF
    try:
        loader = PyPDFLoader(temp_path)
        pages = loader.load_and_split()

        embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(pages, embeddings)

        retriever = vectorstore.as_retriever()
        llm = Ollama(model=model_name)
        qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

        # Summarize button
        if st.button("Summarize RFP"):
            summary_prompt = "Provide a concise summary of the entire RFP document. Focus on the purpose, key requirements, timelines, and submission instructions."
            summary = qa_chain.run(summary_prompt)
            st.session_state.chat_history.append(("Summarize the RFP", summary))
            st.markdown(f"**Summary:**\n\n{summary}")

        # Ask a question
        query = st.text_input("Ask a question about the RFP document:")
        if st.button("Ask"):
            if query:
                answer = qa_chain.run(query)
                st.session_state.chat_history.append((query, answer))
                st.markdown(f"**Answer:** {answer}")
            else:
                st.warning("Please enter a question.")

        # Display chat history
        if st.session_state.chat_history:
            st.subheader("🗂️ Chat History")
            for q, a in st.session_state.chat_history[::-1]:
                st.markdown(f"**Q:** {q}")
                st.markdown(f"**A:** {a}")

        # Export to Word
        if st.button("Export to Word"):
            doc = Document()
            doc.add_heading("RFP Chatbot Conversation", 0)
            for q, a in st.session_state.chat_history:
                doc.add_paragraph(f"Q: {q}", style='List Bullet')
                doc.add_paragraph(f"A: {a}")
            output_path = "chat_history.docx"
            doc.save(output_path)
            with open(output_path, "rb") as f:
                st.download_button("📥 Download Chat History", f, file_name="RFP_Chat_History.docx")

    except Exception as e:
        st.error(f"Something went wrong while processing the PDF: {e}")
